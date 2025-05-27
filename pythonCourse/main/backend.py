import json
import copy

from django.http import JsonResponse, HttpResponse
from django.views.decorators.csrf import csrf_exempt, csrf_protect
from django.views.decorators.http import require_POST
from django.core.files.storage import FileSystemStorage

from .models import User, Statement, Period, Log
from io import BytesIO



# Папка для сохранения файлов
BASE_URL = 'http://127.0.0.1:8000/files/'  # Базовый URL для файлов

########################################
#                  ВСЕ
########################################

# Вход в систему
@csrf_exempt
def sign_in(request):
    try:
        data = json.loads(request.body)
        user = User.get_by_login_password(data.get("login"), data.get("password"))
        if user is not None:
            response = {"answer": user.generate_session()}
        else:
            response = {"answer": False}
    except:
        return HttpResponse("bad request")

    return JsonResponse(response)


# Получить дату начала и конца сбора заявлений
@csrf_exempt
def get_period(request):
    try:
        date = Period.get_period()
        response = {
            "date_start": date[0],
            "date_end": date[1]
        }
    except:
        return HttpResponse("bad request")

    return JsonResponse(response)


# Получить роль
@csrf_exempt
def get_role(request):
    try:
        data = json.loads(request.body)
        user = User.get_by_session(data.get("session"))
        response = {
            "answer": user.get_role()
        }
    except:
        return HttpResponse("bad request")

    return JsonResponse(response)


# Просмотреть заявление
@csrf_exempt
def get_statement(request):
    try:
        data = json.loads(request.body)
        user = User.get_by_session(data.get("session"))
        if user is not None and user.is_student():
            statement = Statement.get_by_id(data.get("statement-id"))

            response = {"statement-data": statement.get_data(),
                        "statement-comments": statement.get_comments(),
                        "statement-json": statement.get_json_data()
                        }
        elif user is not None:
            statement = Statement.get_by_id(data.get("statement-id"))

            response = {"statement-data": statement.get_data(),
                        "statement-json": statement.get_json_data()
                        }
        else:
            response = {"answer": False}
    except:
        return HttpResponse("bad request")

    return JsonResponse(response)


@csrf_exempt
def get_statements(request):
    try:
        data = json.loads(request.body)
        user = User.get_by_session(data.get("session"))
        if user is not None and user.is_student():
            statements = user.get_statements()
            statements_info = []
            for s in statements:
                statements_info.append(s.get_data())
            response = {"statements": statements_info}

        elif user is not None and user.is_inspector():
            response = {"statements": Statement.get_statements_by_statuses(["process", "error", "verified"])}
        elif user is not None and user.is_jury():
            response = {"statements": Statement.get_statements_by_statuses(["conflict", "confirm", "deny", "verified"])}
        elif user is not None and user.is_admin():
            response = {"statements": Statement.get_statements_by_statuses(["process", "error", "verified", "deny", "confirm", "conflict"])}
        else:
            response = {"answer": False}
    except:
        return HttpResponse("bad request")

    return JsonResponse(response)


########################################
#              СТУДЕНТ
########################################



# Просмотреть заявление
@csrf_exempt
def get_info(request):
    try:
        data = json.loads(request.body)
        user = User.get_by_session(data.get("session"))
        if user is not None and user.is_student():
            response = {"FIO": str(user),
                        "phone": user.phone,
                        "email": user.email,
                        "inst": user.inst,
                        "group": user.group,
                        }
        else:
            response = {"answer": False}
    except:
        return HttpResponse("bad request")

    return JsonResponse(response)


# Отправка заявления на проверку
# @csrf_protect
@csrf_exempt
@require_POST
def upload_statement(request):
    try:
        data = request.POST
        json_data_ = data.get("json")
        json_data = json.loads(json_data_)
        user = User.get_by_session(data.get("session"))
        if user is not None and user.is_student() and json_data.get("studies") is not None and json_data.get("science") is not None and json_data.get("activities") is not None and json_data.get("culture") is not None and json_data.get("sport") is not None and (json_data.get("studies") != {} or json_data.get("science") != {} or json_data.get("activities") != {} or json_data.get("culture") != {} or json_data.get("sport") != {}):
            if Statement.upload(user, json.loads(data.get("old_urls")), json_data, request.FILES):
                Log.add(user, "Загрузка заявления", "", copy.deepcopy(Statement.get_by_user(user).get_json_data()))
                response = {"answer": True}
            else:
                response = {"answer": "Too late"}
        else:
            response = {"answer": False}
    except:
        return HttpResponse("bad request")

    #    except Exception as e:
    #    return HttpResponse(str(e))

    return JsonResponse(response)



########################################
#              ИНСПЕКТОР
########################################



# Просмотреть часть заявления
@csrf_exempt
def get_statement_inspector(request):
    try:
        data = json.loads(request.body)
        user = User.get_by_session(data.get("session"))
        if user is not None and user.is_inspector():
            statement_info = Statement.get_part_statement(user, data.get("statement-id"))

            response = {
                "statement": statement_info,
                "type": user.get_role()
            }
        else:
            response = {"answer": False}
    except:
        return HttpResponse("bad request")

    return JsonResponse(response)



# Просмотреть часть заявления
@csrf_exempt
def rate_statement_inspector(request):
    try:
        data = json.loads(request.body)
        user = User.get_by_session(data.get("session"))
        if user is not None and user.is_inspector():
            statement = Statement.get_by_id(data.get("statement-id"))
            statement.mark(user, data.get("mark"), data.get("comment"))
            Log.add(user, "Оценка заявления", "Оценка:" + str(data.get("mark")) + "\nКомментарий:" + data.get("comment"), copy.deepcopy(statement.get_json_data()))
            response = {"answer": True}
        else:
            response = {"answer": False}
    except:
        return HttpResponse("bad request")

    return JsonResponse(response)






########################################
#                 ЖЮРИ
########################################




# Просмотреть часть заявления
@csrf_exempt
def try_statements(request):
    try:
        data = json.loads(request.body)
        user = User.get_by_session(data.get("session"))
        if user is not None and (user.is_jury() or user.is_admin()):
            Statement.system_checkout(int(data.get("count")))
            Log.add(user, "Распределение заявлений, выбрано:" + str(data.get("count")), "", {})
            response = {"answer": True}
        else:
            response = {"answer": False}
    except:
        return HttpResponse("bad request")

    return JsonResponse(response)

# Просмотреть заявление
@csrf_exempt
def get_statement_jury(request):
    try:
        data = json.loads(request.body)
        user = User.get_by_session(data.get("session"))
        if user is not None and user.is_jury():
            statement = Statement.get_by_id(data.get("statement-id"))

            response = {"statement-data": statement.get_data(),
                        "statement-json": statement.get_json_data()
                        }
        else:
            response = {"answer": False}
    except:
        return HttpResponse("bad request")

    return JsonResponse(response)



# Одобрить заявление
@csrf_exempt
def confirm_statement_jury(request):
    try:
        data = json.loads(request.body)
        user = User.get_by_session(data.get("session"))
        if user is not None and user.is_jury():
            statement = Statement.get_by_id(data.get("statement-id"))
            statement.set_status(4)
            Log.add(user, "Одобрение заявления", "", copy.deepcopy(statement.get_json_data()))
            response = {"answer": True}
        else:
            response = {"answer": False}
    except:
        return HttpResponse("bad request")

    return JsonResponse(response)


# Одобрить заявление
@csrf_exempt
def deny_statement_jury(request):
    try:
        data = json.loads(request.body)
        user = User.get_by_session(data.get("session"))
        if user is not None and user.is_jury():
            statement = Statement.get_by_id(data.get("statement-id"))
            statement.set_status(5)
            Log.add(user, "Отклонение заявления", "", copy.deepcopy(statement.get_json_data()))
            response = {"answer": True}
        else:
            response = {"answer": False}
    except:
        return HttpResponse("bad request")

    return JsonResponse(response)


# Выгрузить список человек получивших стипендию
@csrf_exempt
def get_word_success(request):
    try:
        data = json.loads(request.body)
        user = User.get_by_session(data.get("session"))

        if user is not None and (user.is_jury() or user.is_admin()):
            statements = Statement.objects.filter(status__name="confirm", old_status=False)
            temp = []

            for s in statements:
                status = "Спорное"
                if s.status.name == "confirm":
                    status = "Одобрено"
                elif s.status.name == "deny":
                    status = "Отклонено"
                result = f''' {s.user} — {s.points} баллов ({status})
        
        Категории:
            Учёба: {s.mark_studies} ({s.comment_studies})
            Наука: {s.mark_science} ({s.comment_science})
            Мероприятия: {s.mark_activities} ({s.comment_activities})
            Культура: {s.mark_culture} ({s.comment_culture})
            Спорт: {s.mark_sport} ({s.comment_sport})
        Информация для связи:
            Телефон: {s.user.phone}
            Институт: {s.user.inst}
            Группа: {s.user.group}
            Почта: {s.user.email}'''
                temp.append(result)

            Log.add(user, "Выгрузка списка человек получивших стипендию", f"Выполнил: {user}", {})

            buffer, filename = create_word_doc(temp, "Отчет_со_списком_получивших_стипендию")

            response = HttpResponse(buffer.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = f'attachment; filename="{filename}"'
            return response
        else:
            return JsonResponse({"answer": False})
    except Exception as e:
        return HttpResponse(f"bad request: {e}")


# Выгрузить список человек подавших заявления
@csrf_exempt
def get_word_all(request):
    try:
        data = json.loads(request.body)
        user = User.get_by_session(data.get("session"))
        if user is not None and (user.is_jury() or user.is_admin()):
            statements = Statement.objects.filter(old_status=False)
            temp = []

            for s in statements:
                status = "Спорное"
                if s.status.name == "confirm":
                    status = "Одобрено"
                elif s.status.name == "deny":
                    status = "Отклонено"
                result = f''' {s.user} — {s.points} баллов ({status})

            Категории:
                Учёба: {s.mark_studies} ({s.comment_studies})
                Наука: {s.mark_science} ({s.comment_science})
                Мероприятия: {s.mark_activities} ({s.comment_activities})
                Культура: {s.mark_culture} ({s.comment_culture})
                Спорт: {s.mark_sport} ({s.comment_sport})
            Информация для связи:
                Телефон: {s.user.phone}
                Институт: {s.user.inst}
                Группа: {s.user.group}
                Почта: {s.user.email}'''
                temp.append(result)
            Log.add(user, "Выгрузка списка человек получивших стипендию", "Выполнил: "+str(user), {})

            buffer, filename = create_word_doc(temp, "Отчет_со_списком_подавших_заявления")
            response = HttpResponse(buffer.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = f'attachment; filename="{filename}"'
            return response
        else:
            response = {"answer": False}
    except:
        return HttpResponse("bad request")

    return JsonResponse(response)




########################################
#                АДМИН
########################################

# Просмотреть пользователя
@csrf_exempt
def get_user(request):
    try:
        data = json.loads(request.body)
        user = User.get_by_session(data.get("session"))
        if user is not None and user.is_admin():
            user = User.get_by_id(data.get("user-id"))

            response = {
                "user": user.get_data_full(),
            }
        else:
            response = {"answer": False}
    except:
        return HttpResponse("bad request")

    return JsonResponse(response)



# Сменить роль
@csrf_exempt
def change_role_user(request):
    try:
        data = json.loads(request.body)
        user = User.get_by_session(data.get("session"))
        if user is not None and user.is_admin():
            user_ = User.get_by_id(data.get("user-id"))
            user_.change_role(data.get("user-role"))
            Log.add(user, "Изменение роли", data.get("user-id") + " изменен на " + str(data.get("user-role")), {})
            response = {"answer": True}
        else:
            response = {"answer": False}
    except:
        return HttpResponse("bad request")

    return JsonResponse(response)

# Сменить статус заявления
@csrf_exempt
def change_status(request):
    try:
        data = json.loads(request.body)
        user = User.get_by_session(data.get("session"))
        if user is not None and user.is_admin():
            statement = Statement.get_by_id(data.get("statement-id"))
            statement.set_status(data.get("status"))
            Log.add(user, "Изменение статуса заявления", "У заявления " + str(data.get("statment-id")) + " статус изменен на " + str(data.get("status")), {})
            response = {"answer": True}
        else:
            response = {"answer": False}
    except:
        return HttpResponse("bad request")

    return JsonResponse(response)

# Просмотреть список пользователей
@csrf_exempt
def get_list_users(request):
    try:
        data = json.loads(request.body)
        user = User.get_by_session(data.get("session"))
        if user is not None and user.is_admin():

            response = {"users": User.get_users_info()}
        else:
            response = {"answer": False}
    except:
        return HttpResponse("bad request")

    return JsonResponse(response)


# Задать дату начала и конца сбора заявлений
@csrf_exempt
def set_period(request):
    try:
        data = json.loads(request.body)
        user = User.get_by_session(data.get("session"))
        if user is not None and user.is_admin():
            Period.set_start_and_end(data.get("date_start"), data.get("date_end"))
            Log.add(user, "Изменение даты сборов", data.get("date_start") + " --- " + data.get("date_end"), {})
            response = {"answer": True}
        else:
            response = {"answer": False}
    except:
        return HttpResponse("bad request")

    return JsonResponse(response)


# Закешировать все
@csrf_exempt
def set_old(request):
    try:
        data = json.loads(request.body)
        user = User.get_by_session(data.get("session"))
        if user is not None and user.is_admin():
            Statement.set_old()
            Log.add(user, "Перевод заявлений в статус устаревших", "", {})
            response = {"answer": True}
        else:
            response = {"answer": False}
    except:
        return HttpResponse("bad request")

    return JsonResponse(response)


# Просмотреть кешированные
@csrf_exempt
def get_statements_old(request):
    try:
        data = json.loads(request.body)
        user = User.get_by_session(data.get("session"))
        if user is not None and user.is_admin():
            response = {"statements": Statement.get_statements_old()}
        else:
            response = {"answer": False}
    except:
        return HttpResponse("bad request")

    return JsonResponse(response)




# Просмотреть логи
@csrf_exempt
def get_log(request):
    try:
        data = json.loads(request.body)
        user = User.get_by_session(data.get("session"))
        if user is not None and user.is_admin():
            response = {"logs": Log.get_list()}
        else:
            response = {"answer": False}
    except:
        return HttpResponse("bad request")

    return JsonResponse(response)



# Очистить логи
@csrf_exempt
def reset_log(request):
    try:
        data = json.loads(request.body)
        user = User.get_by_session(data.get("session"))
        if user is not None and user.is_admin():
            Log.reset()
            Log.add(user, "Очистка кеша", "", {})
            response = {"answer": True}
        else:
            response = {"answer": False}
    except:
        return HttpResponse("bad request")

    return JsonResponse(response)


# Авто распределение баллов

@csrf_exempt
def auto_points(request):
    try:
        data = json.loads(request.body)
        user = User.get_by_session(data.get("session"))
        if user is not None and user.is_inspector():
            points = 0
            blocks = data.get("blocks")
            points_to_blocks = [
                [75, 70],  # frame-mid (А) получение студентом в течение не менее 2-х последних семестров только оценок
                [25, 20, 10], # frame-mid0 (Б) получение студентом в течение предшествующего года награды (приза) за результаты проектной деятельности и (или) опытно-конструкторской работы
                [25, 20], # frame-mid1 (В) признание студента победителем или призером международной, всероссийской, ведомственной или региональной олимпиады, конкурса, соревнования, состязания или иного мероприятия, направленные на выявление учебных достижений студентов, проведенных в  течение года, предшествующего назначению повышенной государственной академической стипендии
                [12, 8, 5], # frame-mid2 (А.1) Получение наград (призов) за результаты научно-исследовательской работы в рамках мероприятий, организаторами которых являются Министерство науки и высшего образования РФ и государственные организации высшего образования
                [40, 30, 10, 20, 40, 80, 100, 5, 5], # frame-mid3 (А.2) Получение благодарственных писем, грамот органов власти и управления за участие в научных мероприятиях
                [80, 0], # frame-mid4 (А.3) Получение документов, подтверждающих исключительное право студента на достигнутый им научный результат интеллектуальной деятельности, получивших правовую охрану
                [100, 80],  # frame-mid5 (А.4) Получение грантов на выполнение научно-исследовательской работы
                [100, 90, 80, 40, 30, 10, 10],  # frame-mid6 (А.4) (Б) Наличие публикации в научном издании
                [5], # frame-mid7 Критерий «а1». Получение благодарности, благодарственного письма, почетной грамоты за результаты общественной деятельности
                [7],  # frame-mid8
                [10],  # frame-mid9
                [10, 20, 25, 30, 35],  # frame-mid10
                [5, 5, 5, 10, 25],  # frame-mid11 Критерий «а2». Руководство/членство в организациях и объединениях
                [5, 20], # frame-mid12 (А.3) Критерий «а3». Получение грантовой поддержки на реализацию социального проекта
                [1],  # frame-mid13 (А.4) Критерий «а4». Участие в организации и проведении мероприятий работы
                [5], #frame-mid13 (А.4) Критерий «а4». Менеджмент мероприятия
                [7, 10, 20], # frame-mid15 Критерий «а5». Участие в региональных и федеральных конкурсах общественной направленности
                [5], # frame-mid16 Критерий «б». Систематическое участие студента по информационному обеспечению общественно значимых мероприятий, общественной жизни федеральной государственной образовательной организации высшего образования или с её участием, подтверждаемое документально
                [5, 10, 10, 15],  # frame-mid17
                [2], # frame-mid18 Критерий «в». Систематическое участие студента в течение года, предшествующего назначению повышенной государственной академической стипендии, в проведении (обеспечении проведения) общественно значимой деятельности, направленной на формирование у детей и молодежи общероссийской гражданской идентичности, патриотизма и гражданской ответственности, культуры межнационального (межэтнического) и межконфессионального общения, организуемой субъектами, осуществляющими деятельность в сфере молодежной политики, подтверждаемое документально
                [35, 25, 20, 15, 15, 10, 8], # frame-mid19 Критерий «а». Получение студентом в течение года, да, предшествующего назначению повышенной государственной академической стипендии, награды (приза) за результаты культурно-творческой деятельности, осуществленной им в рамках деятельности, проводимой федеральной государственной организацией ВО или иной организацией, в том числе в рамках конкурса, смотра и иного аналогичного международного, всероссийского,  ведомственного, регионального мероприятия, подтверждаемое документально
                [5],  # frame-mid20
                [3], # frame-mid21 Критерий «б». Публичное представление студентом в течение года, предшествующего назнаению повышенной государственной академической стипендии, созданного им произведения литературы или искусства, подтверждаемое документально(в рамках внеучебной деятельности)
                [1], # frame-mid22 Критерий «В1». Систематическое участие студента в течение года, предшествующего назначению повышенной государственной академической стипендии, в проведении (обеспечении проведения) публичной культурно-творческой деятельности воспитательного, пропагандистского характера и иной общественно значимой публично культурно-творческой деятельности, подтверждаемое документально (в рамках внеучебной деятельности)
                [10, 10, 10, 10, 10], # frame-mid23 Критерий "В2" Членство студента в течение года, предшествующего назначению повышенной государственной академической стипендии, в художественных студиях ЦСКиТ
                [50, 25], # frame-mid24 (А.1) В рамках спортивных соревнований международного уровня (олимпийский вид спорта) - I-III место
                [30, 15], # frame-mid25 (А.2) В рамках спортивных соревнований международного уровня (неолимпийский вид спорта) - I-III место
                [50, 25], # frame-mid26 (А.3) В составе сборной команды университета в рамках финальных этапов Всероссийских студенческих соревнований по видам спорта, Чемпионатов (Кубков) студенческих спортивных лиг (ассоциаций) и/или Всероссийской летней/зимней Универсиады - I-III место
                [40, 10], # frame-mid27 (А.4) В рамках спортивных соревнований всероссийского или межрегионального уровня (олимпийский вид спорта) - I-III место
                [20, 10], # frame-mid28 (А.5) В рамках спортивных соревнований всероссийского или межрегионального уровня (не олимпийский вид спорта) - I-III место
                [20, 10], # frame-mid29 (А.6) В рамках спортивных соревнований регионального уровня (олимпийский вид спорта) - I-III место
                [10, 5], # frame-mid30 (А.7) В рамках спортивных соревнований регионального уровня (неолимпийский вид спорта) - I-III место
                [15, 7.5], # frame-mid31 (А.8) В рамках спортивных соревнований муниципального уровня (олимпийский вид спорта) - I-III место
                [10, 5], # frame-mid32 (А.9) В рамках спортивных соревнований муниципального уровня (неолимпийский вид спорта) - I-III место
                [12, 6.5], # frame-mid33 (А.10) В рамках спортивных соревнований, проводимых в университете - I-III место
                [8, 3.5], # frame-mid34 (А.11) В рамках массовых физкультурно-спортивных мероприятий (Кросс нации, Лыжня России, российский Азимут и другие) - I-III
                [30, 15], # frame-mid35 (Б.1) Участие в первенстве/чемпионате международного уровня (олимпийский вид спорта)
                [20, 10], # frame-mid36  (Б.2) Участие в первенстве/чемпионате международного уровня (неолимпийский вид спорта)
                [30, 15], # frame-mid37 (Б.3) Участие в составе сборной команды университета в финальных этапах Всероссийских студенческих соревнований по видам спорта, Чемпионатов (Кубков) студенческих спортивных лиг (ассоциаций) и/или Всероссийской летней/зимней Универсиады
                [20, 10], # frame-mid38 (Б.4) Участие в соревнованиях всероссийского или межрегионального уровня (олимпийский вид спорта)
                [15, 7.5], # frame-mid39 (Б.5) Участие в соревнованиях всероссийского или межрегионального уровня (неолимпийский вид спорта)
                [15, 7.5],  # frame-mid40 (Б.6) Участие в первенстве/чемпионате области (олимпийский вид спорта)
                [10, 5],  # frame-mid41 (Б.7) Участие в первенстве/чемпионате области (неолимпийский вид спорта)
                [10, 5],  # frame-mid42 (Б.8) Участие в первенстве/чемпионате города (олимпийский вид спорта)
                [6, 4.5],  # frame-mid43 (Б.9) Участие в первенстве/чемпионате города (неолимпийский вид спорта)
                [8, 3.5],  # frame-mid44 (Б.10) Участие в спортивных мероприятиях, проводимых в университете
                [5, 2.5],  # frame-mid45 (Б.11) Участие в массовых физкультурно-спортивных мероприятиях (Кросс нации, Лыжня России, Российский Азимут и другие)типо за 2 уже меньше дают баллов
                [8, 7],  # frame-mid46 (Б.12) Судейство соревнований областных, городских, университетских
                [8, 7],  # frame-mid47 (Б.13) Волонтерская деятельность в сфере физической культуры и спорта
                [30],  # frame-mid48 (Б.14) Участие в мероприятиях ВФСК ГТО с получением серебряного знака отличия
                [25],  # frame-mid49 (Б.15) Участие в мероприятиях ВФСК ГТО с получением бронзового знака отличия
                [35] # frame-mid50 (В.1) Выполнение нормативов и требований ВФСК ГТО в официальном центре тестирования, получение золотого знака отличия
            ]
            for i in blocks:
                if i[0] != -1 and i[1] != -1:
                    points += points_to_blocks[i[0]][i[1]]
            response = {"answer": points}
        else:
            response = {"answer": False}
    except:
        return HttpResponse("bad request")

    return JsonResponse(response)




from docx import Document
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement  # 💡 Не забудь этот импорт
from docx.oxml.ns import qn
from datetime import datetime
from io import BytesIO  # 💡 Чтобы возвращать файл как буфер


def create_word_doc(items, title):
    def add_horizontal_line(doc):
        paragraph = doc.add_paragraph()
        p = paragraph._p
        pPr = p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        pPr.append(pBdr)

        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '12')        # потолще линия
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'auto')
        pBdr.append(bottom)

    doc = Document()
    timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    filename = f"{title}_{timestamp}.docx"

    for i, item in enumerate(items, 1):
        para = doc.add_paragraph(style='Normal')
        para.paragraph_format.left_indent = Inches(0)
        para.paragraph_format.first_line_indent = Inches(0)
        # Добавляем настройку интервалов
        para.paragraph_format.space_before = Pt(0)  # Интервал перед параграфом
        para.paragraph_format.space_after = Pt(0)  # Интервал после параграфа

        run = para.add_run(f"{i}) {str(item)}")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)

        add_horizontal_line(doc)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer, filename













